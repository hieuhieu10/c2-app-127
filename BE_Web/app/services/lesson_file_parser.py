from __future__ import annotations

import io
from pathlib import Path


class LessonFileParseError(ValueError):
    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code
        self.message = message


SUPPORTED_LESSON_EXTENSIONS = {".pdf", ".txt", ".docx"}
LEGACY_DOC_EXTENSION = ".doc"


def extract_lesson_text(content: bytes, extension: str) -> str:
    ext = extension.lower()
    if ext == ".txt":
        return _extract_txt(content)
    if ext == ".pdf":
        return _extract_pdf(content)
    if ext == ".docx":
        return _extract_docx(content)
    if ext == LEGACY_DOC_EXTENSION:
        raise LessonFileParseError(
            "LEGACY_DOC_UNSUPPORTED",
            "File .doc cũ chưa được hỗ trợ ổn định. Vui lòng lưu giáo án thành .docx, .pdf hoặc .txt.",
        )
    raise LessonFileParseError("UNSUPPORTED_FILE_TYPE", "Chỉ hỗ trợ file PDF, TXT hoặc DOCX.")


def build_preview(text: str, limit: int = 500) -> str:
    compact = " ".join(text.split())
    return compact[:limit]


def truncate_source_text(text: str, max_chars: int) -> str:
    normalized = text.strip()
    if len(normalized) <= max_chars:
        return normalized
    return normalized[:max_chars].rstrip()


def _extract_txt(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "cp1258", "cp1252"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace").strip()


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise LessonFileParseError(
            "PARSER_DEPENDENCY_MISSING",
            "Backend chưa cài pypdf để đọc PDF. Chạy uv sync hoặc pip install pypdf.",
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(content))
        pages = [
            f"===== PAGE {index} =====\n{page_text.strip()}"
            for index, page in enumerate(reader.pages, start=1)
            if (page_text := (page.extract_text() or "")).strip()
        ]
    except Exception as exc:  # pragma: no cover - pypdf raises varied parser exceptions
        raise LessonFileParseError("PARSE_FAILED", "Không đọc được nội dung PDF.") from exc
    return "\n\n".join(pages).strip()


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise LessonFileParseError(
            "PARSER_DEPENDENCY_MISSING",
            "Backend chưa cài python-docx để đọc DOCX. Chạy uv sync hoặc pip install python-docx.",
        ) from exc

    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:  # pragma: no cover - python-docx raises varied parser exceptions
        raise LessonFileParseError("PARSE_FAILED", "Không đọc được nội dung DOCX.") from exc

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    table_cells.append(text)
    return "\n".join([*paragraphs, *table_cells]).strip()


def safe_extension(filename: str) -> str:
    return Path(filename or "").suffix.lower()
